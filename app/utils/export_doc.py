import io
import json
import os
from pathlib import Path
import traceback
from typing import Any, Dict, Optional

import chardet  # Thêm thư viện để phát hiện encoding
from docx import Document
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import JSONResponse

from app.storage import pgdb
from app.utils.md_to_docx import convert_markdown_to_docx

# Đường dẫn đến thư mục Downloads
DOWNLOADS_DIR = str(Path.home() / "Downloads")

# Đường dẫn đến file template mặc định - cập nhật theo file mới
TEMPLATE_PATH = "TemplateTBDU_Bo22.docx"
BASE_DIR = os.environ.get('APP_BASE_DIR', os.path.dirname(
    os.path.dirname(os.path.dirname(__file__))))
router = APIRouter(
    prefix="/export_doc",
    tags=["export_doc"],
    dependencies=[],
    responses={404: {"description": "Not found"}},
)
TEMPDIR = os.path.join(BASE_DIR, "temp")

def set_table_border_style(table):
    """
    Thiết lập viền nét liền cho toàn bộ bảng
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # Đặt thuộc tính căn chỉnh cho bảng
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Lặp qua tất cả các ô trong bảng
    for row in table.rows:
        for cell in row.cells:
            # Lấy thuộc tính ô
            tcPr = cell._element.tcPr

            # Tạo đường viền cho ô
            tcBorders = OxmlElement('w:tcBorders')

            # Định nghĩa các cạnh viền (top, left, bottom, right)
            for border_position in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_position}')
                border.set(qn('w:val'), 'single')  # Đường nét liền
                border.set(qn('w:sz'), '4')  # Độ dày (4 = 1pt)
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')  # Màu đen
                tcBorders.append(border)

            # Nếu đã có tcBorders, xóa nó trước
            existing_borders = tcPr.find(qn('w:tcBorders'))
            if existing_borders is not None:
                tcPr.remove(existing_borders)

            # Thêm tcBorders mới vào
            tcPr.append(tcBorders)

    return table


def fill_table_with_json_data(doc: Document, json_data: Dict[Any, Any]) -> Document:
    """
    Điền dữ liệu JSON vào bảng của file DOCX mẫu
    """
    # Lấy bảng đầu tiên từ document
    if len(doc.tables) == 0:
        raise ValueError("Template document does not contain any tables")

    table = doc.tables[0]

    # Lấy dữ liệu từ JSON - điều chỉnh theo cấu trúc của bo22.txt
    requirement_level_0 = json_data.get("requirement_level_0", {})

    # Tìm hàng bắt đầu để điền dữ liệu (hàng đầu tiên không phải header)
    start_row = 0
    for i, row in enumerate(table.rows):
        # Kiểm tra nếu đây là hàng header
        if any(cell.text.strip().startswith("Mục") for cell in row.cells):
            start_row = i + 1  # Bắt đầu điền dữ liệu từ hàng sau header
            break

    # Xác định vị trí các cột trong bảng
    header_row = table.rows[start_row - 1] if start_row > 0 else table.rows[0]
    col_indices = {}
    for i, cell in enumerate(header_row.cells):
        cell_text = cell.text.strip().lower()
        if "mục" in cell_text:
            col_indices["muc"] = i
        elif "yêu cầu hàng hóa" in cell_text or "thông số kỹ thuật tối thiểu" in cell_text:
            col_indices["requirement"] = i
        elif "thông số kỹ thuật/yêu cầu dịch vụ" in cell_text:
            col_indices["specification"] = i
        elif "mức độ đáp ứng" in cell_text:
            col_indices["compliance"] = i
        elif "tham chiếu" in cell_text or "ghi chú" in cell_text:
            col_indices["reference"] = i

    # Sử dụng vị trí mặc định nếu không tìm thấy
    if "muc" not in col_indices:
        col_indices["muc"] = 0
    if "requirement" not in col_indices:
        col_indices["requirement"] = 1
    if "specification" not in col_indices:
        col_indices["specification"] = 2
    if "compliance" not in col_indices:
        col_indices["compliance"] = 3
    if "reference" not in col_indices:
        col_indices["reference"] = 4

    # Current row index to insert data
    row_index = start_row

    # Lưu trữ nội dung cột "Yêu cầu hàng hóa" cho các cấp
    requirement_texts = {
        0: {},  # Sẽ được điền động theo mục
        1: {},  # Sẽ được điền động theo mục
        2: {},  # Sẽ được điền động theo mục
        3: {}   # Sẽ được điền động theo mục
    }

    # Hàm đệ quy để xử lý và điền các yêu cầu vào bảng
    def process_requirements(requirement, level=0, parent_muc=""):
        nonlocal row_index

        if level == 0:
            # Process level 0 (main requirement)
            muc = requirement.get("muc", "")
            req_name = requirement.get("requirement_name", "")

            # Lưu thông tin yêu cầu cấp 0
            requirement_texts[0][muc] = req_name

            # Kiểm tra xem có đủ hàng không, nếu không thì tạo thêm
            while row_index >= len(table.rows):
                table.add_row()

            # Điền dữ liệu vào hàng hiện tại
            row = table.rows[row_index]
            row.cells[col_indices["muc"]].text = muc
            row.cells[col_indices["requirement"]].text = req_name

            # Kiểm tra nếu có description trực tiếp ở level 0
            if "description" in requirement:
                descriptions = requirement.get("description", [])
                if descriptions:
                    for i, desc in enumerate(descriptions):
                        desc_detail = desc.get("description_detail", "")
                        # Điền vào cột specification của hàng hiện tại nếu là mô tả đầu tiên
                        if i == 0 and col_indices["specification"] < len(row.cells):
                            row.cells[col_indices["specification"]
                                      ].text = desc_detail
                        # Hoặc tạo hàng mới cho các mô tả tiếp theo
                        else:
                            row_index += 1
                            while row_index >= len(table.rows):
                                table.add_row()
                            row = table.rows[row_index]
                            if col_indices["specification"] < len(row.cells):
                                row.cells[col_indices["specification"]
                                          ].text = desc_detail

            row_index += 1

            # Process sub-requirements
            sub_reqs = requirement.get("sub_requirements", [])
            for sub_req in sub_reqs:
                for key, value in sub_req.items():
                    process_requirements(value, 1, muc)

        elif level == 1:
            # Process level 1
            muc = requirement.get("muc", "")
            req_name = requirement.get("requirement_name", "")

            # Lưu thông tin yêu cầu cấp 1
            requirement_texts[1][muc] = req_name

            # Kiểm tra xem có đủ hàng không, nếu không thì tạo thêm
            while row_index >= len(table.rows):
                table.add_row()

            # Điền dữ liệu vào hàng hiện tại
            row = table.rows[row_index]
            row.cells[col_indices["muc"]].text = muc
            row.cells[col_indices["requirement"]].text = req_name

            # Kiểm tra nếu có description trực tiếp ở level 1
            if "description" in requirement:
                descriptions = requirement.get("description", [])
                if descriptions:
                    for i, desc in enumerate(descriptions):
                        desc_detail = desc.get("description_detail", "")
                        # Điền vào cột specification của hàng hiện tại nếu là mô tả đầu tiên
                        if i == 0 and col_indices["specification"] < len(row.cells):
                            row.cells[col_indices["specification"]
                                      ].text = desc_detail
                        # Hoặc tạo hàng mới cho các mô tả tiếp theo
                        else:
                            row_index += 1
                            while row_index >= len(table.rows):
                                table.add_row()
                            row = table.rows[row_index]
                            if col_indices["specification"] < len(row.cells):
                                row.cells[col_indices["specification"]
                                          ].text = desc_detail

            row_index += 1

            # Process sub-requirements
            sub_reqs = requirement.get("sub_requirements", [])
            for sub_req in sub_reqs:
                for key, value in sub_req.items():
                    process_requirements(value, 2, muc)

        elif level == 2:
            # Process level 2
            muc = requirement.get("muc", "")
            req_name = requirement.get("requirement_name", "")

            # Lưu thông tin yêu cầu cấp 2
            requirement_texts[2][muc] = req_name

            # Tìm parent requirement
            parent_req = requirement_texts[1].get(parent_muc, "")

            # Kiểm tra xem có đủ hàng không, nếu không thì tạo thêm
            while row_index >= len(table.rows):
                table.add_row()

            # Điền dữ liệu vào hàng hiện tại
            row = table.rows[row_index]
            row.cells[col_indices["muc"]].text = muc
            row.cells[col_indices["requirement"]].text = req_name

            # Kiểm tra nếu có description trực tiếp ở level 2
            has_description = False
            if "description" in requirement:
                descriptions = requirement.get("description", [])
                if descriptions:
                    has_description = True
                    for i, desc in enumerate(descriptions):
                        desc_detail = desc.get("description_detail", "")
                        # Điền vào cột specification của hàng hiện tại nếu là mô tả đầu tiên
                        if i == 0 and col_indices["specification"] < len(row.cells):
                            row.cells[col_indices["specification"]
                                      ].text = desc_detail
                        # Hoặc tạo hàng mới cho các mô tả tiếp theo
                        else:
                            row_index += 1
                            while row_index >= len(table.rows):
                                table.add_row()
                            row = table.rows[row_index]
                            if col_indices["specification"] < len(row.cells):
                                row.cells[col_indices["specification"]
                                          ].text = desc_detail

            # Nếu không có description, vẫn tăng row_index lên
            if not has_description:
                row_index += 1
            else:
                # Nếu có description và không phải là mô tả đầu tiên, row_index đã được tăng trong vòng lặp description
                pass

            # Xử lý các sub-requirements ở cấp 3
            sub_reqs = requirement.get("sub_requirements", [])
            if sub_reqs:
                for sub_req in sub_reqs:
                    for key, value in sub_req.items():
                        process_requirements(value, 3, muc)

        elif level == 3:
            # Process level 3 with descriptions
            muc = requirement.get("muc", "")  # Thêm mục ở level 3 nếu có
            req_name = requirement.get("requirement_name", "")
            descriptions = requirement.get("description", [])

            # Lưu thông tin yêu cầu cấp 3 nếu có
            if req_name:
                requirement_texts[3][parent_muc] = req_name

                # Thêm hàng hiển thị tên yêu cầu cấp 3
                while row_index >= len(table.rows):
                    table.add_row()

                row = table.rows[row_index]
                if muc:  # Nếu có mục ở level 3, hiển thị nó
                    row.cells[col_indices["muc"]].text = muc
                row.cells[col_indices["requirement"]].text = req_name

                # Kiểm tra nếu có description và hiển thị description đầu tiên trên cùng một hàng
                if descriptions and len(descriptions) > 0:
                    desc_detail = descriptions[0].get("description_detail", "")
                    if col_indices["specification"] < len(row.cells):
                        row.cells[col_indices["specification"]
                                  ].text = desc_detail

                    # Xử lý các description tiếp theo trên các hàng riêng biệt
                    for i in range(1, len(descriptions)):
                        row_index += 1
                        while row_index >= len(table.rows):
                            table.add_row()

                        row = table.rows[row_index]
                        desc_detail = descriptions[i].get(
                            "description_detail", "")
                        if col_indices["specification"] < len(row.cells):
                            row.cells[col_indices["specification"]
                                      ].text = desc_detail

                row_index += 1
            else:
                # Nếu không có requirement_name, chỉ hiển thị các descriptions
                for desc in descriptions:
                    desc_detail = desc.get("description_detail", "")

                    # Kiểm tra xem có đủ hàng không, nếu không thì tạo thêm
                    while row_index >= len(table.rows):
                        table.add_row()

                    # Điền dữ liệu vào hàng hiện tại
                    row = table.rows[row_index]

                    # Điền thông số kỹ thuật/yêu cầu dịch vụ
                    if col_indices["specification"] < len(row.cells):
                        row.cells[col_indices["specification"]
                                  ].text = desc_detail

                    row_index += 1

    # Xóa dữ liệu hiện có trong bảng (ngoại trừ hàng header)
    for i in range(start_row, len(table.rows)):
        for cell in table.rows[i].cells:
            cell.text = ""

    # Process the main requirement
    process_requirements(requirement_level_0)

    # Thiết lập viền nét liền cho bảng
    table = set_table_border_style(table)

    return doc


async def get_template_doc(template_file: Optional[UploadFile] = None) -> Document:
    """
    Lấy Document từ template file hoặc từ template mặc định
    """
    if template_file:
        # Sử dụng template file từ người dùng
        content = await template_file.read()
        doc = Document(io.BytesIO(content))
    else:
        # Sử dụng template mặc định
        if not os.path.exists(TEMPLATE_PATH):
            raise HTTPException(
                status_code=404, detail=f"Default template file not found: {TEMPLATE_PATH}")
        doc = Document(TEMPLATE_PATH)

    return doc


@router.post("/export_docs/")
async def export_docs(
    page_id: str = "",  # Giữ tham số này để tương thích với signature đã có
    json_content: str = Form(...),
    output_filename: str = Form("response_document.docx"),
    template_file: Optional[UploadFile] = File(None)
):
    """
    Điền dữ liệu JSON vào file DOCX mẫu và lưu vào thư mục Downloads

    Args:
        page_id: Ignored, kept for compatibility
        json_content: JSON string chứa dữ liệu yêu cầu
        output_filename: Tên mong muốn cho file DOCX đầu ra
        template_file: File template DOCX (nếu không cung cấp, sẽ sử dụng mặc định)

    Returns:
        JSON response với thông tin về file
    """
    try:
        # Parse JSON
        try:
            # Đảm bảo json_content là chuỗi trước khi parse
            if isinstance(json_content, bytes):
                # Phát hiện mã hóa của chuỗi bytes
                detection = chardet.detect(json_content)
                detected_encoding = detection['encoding']
                if detected_encoding:
                    json_content = json_content.decode(detected_encoding)
                else:
                    # Thử các encoding phổ biến
                    try:
                        json_content = json_content.decode('utf-8')
                    except UnicodeDecodeError:
                        try:
                            json_content = json_content.decode('utf-16')
                        except UnicodeDecodeError:
                            json_content = json_content.decode('latin-1')

            json_data = json.loads(json_content)
        except json.JSONDecodeError as e:
            raise HTTPException(
                status_code=400, detail=f"Invalid JSON format: {str(e)}")
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error processing JSON content: {str(e)}")

        # Lấy document từ template
        try:
            doc = await get_template_doc(template_file)
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error loading template: {str(e)}")

        # Điền dữ liệu vào document
        try:
            doc = fill_table_with_json_data(doc, json_data)
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error filling template: {str(e)}")

        # Đảm bảo tên file có phần mở rộng .docx
        if not output_filename.lower().endswith('.docx'):
            output_filename += '.docx'

        # Đảm bảo thư mục Downloads tồn tại
        os.makedirs(DOWNLOADS_DIR, exist_ok=True)

        # Tạo đường dẫn đến file
        file_path = os.path.join(DOWNLOADS_DIR, output_filename)

        # Xử lý trùng tên file
        counter = 1
        base_name, ext = os.path.splitext(output_filename)
        while os.path.exists(file_path):
            new_filename = f"{base_name}_{counter}{ext}"
            file_path = os.path.join(DOWNLOADS_DIR, new_filename)
            counter += 1

        # Lưu document trực tiếp vào thư mục Downloads
        doc.save(file_path)

        # Trả về phản hồi thành công với thông tin về file
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": "Document exported successfully",
                "file_path": file_path,
                "file_name": os.path.basename(file_path)
            }
        )

    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Error processing document: {str(e)}")

def export_docs_from_file(
    proposal_id: str = Form(...),
    output_filename: str = Form("Ho_so_ky_thuat.docx")
):
    """
    Điền dữ liệu JSON từ cơ sở dữ liệu vào file DOCX Template.docx từ thư mục temp và lưu vào thư mục Downloads

    Args:
        proposal_id: ID của proposal để lấy dữ liệu từ database
        output_filename: Tên mong muốn cho file DOCX đầu ra

    Returns:
        JSON response với thông tin về file
    """
    try:
        # Lấy dữ liệu JSON từ cơ sở dữ liệu bằng proposal_id

        sql = f"""select * from technical_requirement_json where proposal_id = {proposal_id}"""
        contentdb = pgdb.select(sql)

        # Kiểm tra nếu không có dữ liệu
        if not contentdb or len(contentdb) == 0:
            raise HTTPException(
                status_code=404, detail=f"No data found for proposal_id: {proposal_id}")

        # Lấy trường requirement_json từ kết quả truy vấn
        content_raw = contentdb[0]["requirement_json"]

        # Kiểm tra nếu nội dung trống
        if not content_raw:
            raise HTTPException(
                status_code=400, detail="JSON content from database is empty")

        # Parse JSON data
        try:
            # Nếu đã là dict hoặc list (đã được parse bởi PostgreSQL)
            if isinstance(content_raw, (dict, list)):
                json_data = content_raw
            else:
                # Nếu là chuỗi, cần parse
                # Phát hiện mã hóa nếu là bytes
                if isinstance(content_raw, bytes):
                    detection = chardet.detect(content_raw)
                    detected_encoding = detection['encoding']
                    if detected_encoding:
                        content_str = content_raw.decode(detected_encoding)
                    else:
                        content_str = content_raw.decode('utf-8')
                else:
                    content_str = content_raw

                # Loại bỏ BOM nếu có
                if isinstance(content_str, str) and content_str.startswith('\ufeff'):
                    content_str = content_str[1:]

                # Parse JSON
                json_data = json.loads(content_str)

            # Kiểm tra nếu JSON trống
            if not json_data:
                raise HTTPException(
                    status_code=400, detail="JSON data is empty after parsing")

        except json.JSONDecodeError as e:
            # Hiển thị thông tin lỗi nếu có
            if isinstance(content_raw, str):
                error_pos = e.pos
                snippet = content_raw[max(
                    0, error_pos-20):min(len(content_raw), error_pos+20)]
                details = f"Invalid JSON format at position {error_pos}: {e.msg}. Content around error: '{snippet}'"
            else:
                details = f"Invalid JSON format: {str(e)}"
            raise HTTPException(status_code=400, detail=details) from e
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error processing JSON content: {str(e)}") from e

        # Sử dụng Template.docx từ thư mục temp trong dự án
        try:
            template_path = os.path.join(BASE_DIR,"temp", "Template.docx")

            if not os.path.exists(template_path):
                raise HTTPException(
                    status_code=404, detail=f"Default template file not found: {template_path} at {traceback.format_exc()}")
            doc = Document(template_path)
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error loading template: {str(e)}") from e

        # Điền dữ liệu vào document
        try:
            doc = fill_table_with_json_data(doc, json_data)
        except Exception as e:
            raise HTTPException(
                status_code=500, detail=f"Error filling template: {str(e)}") from e

        # Đảm bảo tên file có phần mở rộng .docx
        if not output_filename.lower().endswith('.docx'):
            output_filename += '.docx'

        # Đảm bảo thư mục temp tồn tại
        os.makedirs(TEMPDIR, exist_ok=True)

        # Tạo đường dẫn đến file
        file_path = os.path.join(TEMPDIR, output_filename)

        # Xử lý trùng tên file
        counter = 1
        base_name, ext = os.path.splitext(output_filename)
        while os.path.exists(file_path):
            new_filename = f"{base_name}_{counter}{ext}"
            file_path = os.path.join(TEMPDIR, new_filename)
            counter += 1

        # Lưu document trực tiếp vào thư mục temp
        doc.save(file_path)

        # Trả về phản hồi thành công với thông tin về file
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": "Document exported successfully to temp folder",
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "proposal_id": proposal_id
            }
        )

    except HTTPException:
        # Đẩy lên lại các HTTP Exception đã xác định
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Error processing document: {str(e)}") from e


@router.post("/debug_json_file/")
async def debug_json_file(
    json_file: UploadFile = File(...)
):
    """
    Kiểm tra và hiển thị thông tin về file JSON để debug

    Args:
        json_file: File TXT chứa dữ liệu JSON

    Returns:
        JSON response với thông tin về nội dung file
    """
    try:
        # Đọc nội dung file
        content = await json_file.read()

        # Thông tin cơ bản về file
        file_info = {
            "file_size": len(content),
            "file_name": json_file.filename,
            "content_type": json_file.content_type
        }

        # Phát hiện encoding
        detection = chardet.detect(content)
        file_info["detected_encoding"] = detection

        # Thử tất cả các encoding phổ biến
        encodings = ["utf-8", "utf-16", "utf-16-le",
                     "utf-16-be", "latin-1", "cp1252"]
        decoded_results = {}

        for enc in encodings:
            try:
                decoded = content.decode(enc)
                # Loại bỏ BOM nếu có
                if decoded.startswith('\ufeff'):
                    decoded = decoded[1:]

                decoded_results[enc] = {
                    "success": True,
                    "preview": decoded[:50],
                    "length": len(decoded)
                }

                # Thử parse nội dung này thành JSON
                try:
                    json_data = json.loads(decoded)
                    decoded_results[enc]["is_valid_json"] = True

                    if isinstance(json_data, dict):
                        decoded_results[enc]["top_keys"] = list(
                            json_data.keys())
                        if "response" in json_data and "requirement_level_0" in json_data.get("response", {}):
                            decoded_results[enc]["has_required_structure"] = True
                    else:
                        decoded_results[enc]["is_dict"] = False

                except json.JSONDecodeError as e:
                    decoded_results[enc]["is_valid_json"] = False
                    decoded_results[enc]["json_error"] = str(e)

            except UnicodeDecodeError as e:
                decoded_results[enc] = {
                    "success": False,
                    "error": str(e)
                }

        # Thêm kết quả giải mã vào thông tin file
        file_info["decoded_results"] = decoded_results

        # Trả về thông tin debug
        return JSONResponse(
            status_code=200,
            content={
                "message": "JSON file analysis results",
                "file_info": file_info
            }
        )

    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={
                "message": f"Error analyzing file: {str(e)}"
            }
        )


def convert_md_to_docx(
    markdown_content: str = Form(...),
    output_filename: str = Form("markdown_document.docx"),
):
    """
    Convert Markdown content to DOCX format

    Args:
        markdown_content: String containing markdown content
        output_filename: Desired name for the output DOCX file
        template_file: Optional template DOCX file

    Returns:
        JSON response with information about the created file
    """
    try:
        # Get document from template or create a new one

        doc = Document()

        # Convert markdown to docx
        doc = convert_markdown_to_docx(doc, markdown_content)

        # Ensure output filename has .docx extension
        if not output_filename.lower().endswith('.docx'):
            output_filename += '.docx'

        # Ensure Downloads directory exists
        os.makedirs(DOWNLOADS_DIR, exist_ok=True)

        # Create path to file
        file_path = os.path.join(DOWNLOADS_DIR, output_filename)

        # Handle duplicate filenames
        counter = 1
        base_name, ext = os.path.splitext(output_filename)
        while os.path.exists(file_path):
            new_filename = f"{base_name}_{counter}{ext}"
            file_path = os.path.join(DOWNLOADS_DIR, new_filename)
            counter += 1

        # Save document to Downloads folder
        doc.save(file_path)

        # Return success response with file info
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": "Markdown converted to DOCX successfully",
                "file_path": file_path,
                "file_name": os.path.basename(file_path)
            }
        )

    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Error converting markdown to DOCX: {str(e)}") from e 
