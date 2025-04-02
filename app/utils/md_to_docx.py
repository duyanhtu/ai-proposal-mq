import re

import markdown
from bs4 import BeautifulSoup
from docx.shared import Inches, Pt, RGBColor


def preprocess_markdown(markdown_text):
    """
    Preprocess markdown text to handle literal newline characters and other formatting issues

    Args:
        markdown_text: Raw markdown text that might contain literal \n sequences

    Returns:
        Cleaned markdown text ready for conversion
    """
    # Remove HTML comments
    processed_text = re.sub(r'<!--.*?-->', '', markdown_text, flags=re.DOTALL)

    # First: handle special case of list items with embedded \n
    # Replace patterns like "- **Item:**\n   - Subitem" keeping the list structure
    processed_text = re.sub(
        r'([-*]\s+\*\*[^*]+\*\*:?)\\n\s*([-*]\s+)', r'\1\n\2', processed_text)

    # Handle numbered lists with embedded \n
    processed_text = re.sub(
        r'(\*\*\d+\.\s+[^*]+\*\*:?)\\n\s*([-*]\s+)', r'\1\n\2', processed_text)

    # Now replace remaining literal \n with actual newlines
    processed_text = re.sub(r'\\n', '\n', processed_text)

    # Fix lists where a newline would break the item
    # Look for list items that end with a colon followed by a newline
    processed_text = re.sub(
        r'([-*]\s+[^:\n]+:)\n(\s+)', r'\1 ', processed_text)

    # Fix inconsistent heading formatting
    processed_text = re.sub(r'\n#+\s*([^\n]+)', r'\n\n# \1', processed_text)

    # Ensure proper spacing for list items
    processed_text = re.sub(r'\n\s*[-*]\s', r'\n\n- ', processed_text)

    # Fix bullet points within a line (common after converting \n)
    processed_text = re.sub(r'(\S)\n\s*[-*]\s+', r'\1\n\n- ', processed_text)

    # Fix multiple consecutive newlines
    processed_text = re.sub(r'\n{3,}', '\n\n', processed_text)

    return processed_text


def convert_markdown_to_docx(doc, markdown_text):
    """
    Convert markdown text to Word document format

    Args:
        doc: Document object from python-docx
        markdown_text: String containing markdown content

    Returns:
        Updated document object with markdown content converted to Word format
    """
    # Preprocess the markdown to handle literal newlines
    cleaned_markdown = preprocess_markdown(markdown_text)

    # Convert markdown to HTML with extensions
    html = markdown.markdown(
        cleaned_markdown,
        extensions=[
            'tables',
            'fenced_code',
            'codehilite',
            'nl2br',  # Convert newlines to <br> tags
            'extra'   # Extra markdown features
        ]
    )

    # Parse the HTML
    soup = BeautifulSoup(html, 'html.parser')

    # Process the entire HTML body - Fixed the None error
    if soup is None:
        # If parsing completely failed, add the text directly
        doc.add_paragraph(cleaned_markdown)
        return doc

    # Use find_all instead of trying to access children directly
    for element in soup.find_all(True, recursive=False):
        process_element(element, doc)

    # Handle case where there might be direct text nodes
    if soup.string and soup.string.strip():
        doc.add_paragraph(soup.string.strip())

    return doc


def process_element(element, doc, list_level=0):
    """Process an HTML element and add it to the Word document"""
    if element is None:
        return

    if element.name is None:
        # Plain text
        if element.string and element.string.strip():
            p = doc.add_paragraph()
            p.add_run(element.string.strip())
        return

    if element.name == 'h1':
        p = doc.add_paragraph(style='Heading 1')
        p.add_run(element.get_text().strip())

    elif element.name == 'h2':
        p = doc.add_paragraph(style='Heading 2')
        p.add_run(element.get_text().strip())

    elif element.name == 'h3':
        p = doc.add_paragraph(style='Heading 3')
        p.add_run(element.get_text().strip())

    elif element.name == 'h4':
        p = doc.add_paragraph(style='Heading 4')
        p.add_run(element.get_text().strip())

    elif element.name in ['h5', 'h6']:
        p = doc.add_paragraph(style='Heading 5')
        p.add_run(element.get_text().strip())

    elif element.name == 'p':
        p = doc.add_paragraph()
        process_inline_content(element, p)

    elif element.name == 'br':
        # For line breaks, add a run with a newline
        if doc.paragraphs:
            p = doc.paragraphs[-1]
        else:
            p = doc.add_paragraph()
        p.add_run('\n')

    elif element.name == 'ul':
        for li in element.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(list_level * 0.25)
            # Process the list item content in a more controlled way
            process_inline_content(li, p, is_list_item=True)

            # Process any nested lists
            for nested_list in li.find_all(['ul', 'ol'], recursive=False):
                process_element(nested_list, doc, list_level + 1)

    elif element.name == 'ol':
        for i, li in enumerate(element.find_all('li', recursive=False), 1):
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(list_level * 0.25)
            # Process the list item content in a more controlled way
            process_inline_content(li, p, is_list_item=True)

            # Process any nested lists
            for nested_list in li.find_all(['ul', 'ol'], recursive=False):
                process_element(nested_list, doc, list_level + 1)

    elif element.name in ['pre', 'code']:
        code_text = element.get_text()
        p = doc.add_paragraph()
        code_run = p.add_run(code_text)
        code_run.font.name = 'Courier New'
        code_run.font.size = Pt(10)
        p.paragraph_format.left_indent = Inches(0.5)

    elif element.name == 'blockquote':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.style = 'Quote'
        process_inline_content(element, p)

    elif element.name == 'hr':
        doc.add_paragraph('_' * 50)

    # Process tables
    elif element.name == 'table':
        rows = element.find_all('tr')
        if rows:
            num_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
            if num_cols > 0:
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'

                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        if j < num_cols:
                            table_cell = table.cell(i, j)
                            process_inline_content(
                                cell, table_cell.paragraphs[0])

    # Handle div and other container elements
    elif element.name in ['div', 'section', 'article']:
        for child in element.contents:
            if isinstance(child, str) and child.strip():
                p = doc.add_paragraph()
                p.add_run(child.strip())
            elif hasattr(child, 'name'):
                process_element(child, doc, list_level)

    # Recursively process children for other elements
    elif element.name not in ['li', 'td', 'th']:
        for child in element.contents:
            if isinstance(child, str) and child.strip():
                p = doc.add_paragraph()
                p.add_run(child.strip())
            elif hasattr(child, 'name'):
                process_element(child, doc, list_level)


def process_inline_content(element, paragraph, is_list_item=False):
    """Process inline elements within a paragraph"""
    if element is None or not hasattr(element, 'contents') or not element.contents:
        return

    for content in element.contents:
        if content is None:
            continue

        if not hasattr(content, 'name'):  # Plain text
            if isinstance(content, str) and content.strip():
                # For list items, we want to keep everything in one line
                if is_list_item:
                    # Replace newlines with spaces in list items
                    text = content.replace('\n', ' ')
                    paragraph.add_run(text)
                else:
                    # Handle potential newlines in text for non-list items
                    text_parts = content.split('\n')
                    for i, part in enumerate(text_parts):
                        if part:
                            paragraph.add_run(part)
                        if i < len(text_parts) - 1:  # Add line break except after last part
                            paragraph.add_run('\n')
        elif content.name is None:  # Plain text as NavigableString
            if content.string:
                # For list items, we want to keep everything in one line
                if is_list_item:
                    # Replace newlines with spaces in list items
                    text = content.string.replace('\n', ' ')
                    paragraph.add_run(text)
                else:
                    # Handle potential newlines in text for non-list items
                    text_parts = content.string.split('\n')
                    for i, part in enumerate(text_parts):
                        if part:
                            paragraph.add_run(part)
                        if i < len(text_parts) - 1:  # Add line break except after last part
                            paragraph.add_run('\n')

        elif content.name == 'strong' or content.name == 'b':
            run = paragraph.add_run(content.get_text())
            run.bold = True
        elif content.name == 'em' or content.name == 'i':
            run = paragraph.add_run(content.get_text())
            run.italic = True
        elif content.name == 'u':
            run = paragraph.add_run(content.get_text())
            run.underline = True
        elif content.name == 'code':
            run = paragraph.add_run(content.get_text())
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(89, 89, 89)
        elif content.name == 'a':
            run = paragraph.add_run(content.get_text())
            run.underline = True
            run.font.color.rgb = RGBColor(0, 0, 255)
        elif content.name == 'br':
            if not is_list_item:  # Skip line breaks in list items
                paragraph.add_run('\n')
        elif content.name in ['ul', 'ol']:
            # Skip lists as they're handled separately
            pass
        else:
            # Recursively process other inline elements
            process_inline_content(content, paragraph, is_list_item)
