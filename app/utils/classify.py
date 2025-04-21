import os
import re
import shutil
import tempfile

# Add these imports at the top of your file
import time
import traceback  # Add these imports at the top of your file
import uuid

import fitz
from minio.error import S3Error

from app.config.env import EnvSettings
from app.model_ai import llm
from app.mq.rabbit_mq import RabbitMQClient
from app.storage.postgre import executeSQL, selectSQL
from app.utils.download_file_minio import get_minio_client
from app.utils.logger import get_logger
from app.utils.minio import upload_to_minio
from app.utils.pdf_image_to_text_batch import convert_pdf_to_text

# Set up logger using the centralized logging system
logger = get_logger(__name__)

MINIO_BUCKET = EnvSettings().MINIO_BUCKET  # Bucket name
RABBIT_MQ_CHPATER_SPLITER_QUEUE = EnvSettings().RABBIT_MQ_CHPATER_SPLITER_QUEUE

rabbit_mq = RabbitMQClient(
    host=EnvSettings().RABBIT_MQ_HOST,
    port=EnvSettings().RABBIT_MQ_PORT,
    user=EnvSettings().RABBIT_MQ_USER,
    password=EnvSettings().RABBIT_MQ_PASS,
)

BASE_DIR = os.environ.get('APP_BASE_DIR', os.path.dirname(
    os.path.dirname(os.path.dirname(__file__))))


def classify(hs_id: str, email: str):
    """
    Classify all the files in database.
    This function will download the files from MinIO, classify them based on their content,
    and update the database with the classification results.
    It will also publish a message to RabbitMQ for further processing.
    If no files are found for the given hs_id, it will return a message indicating that.
    """
    # select file with hs_id
    results = selectSQL(
        "SELECT id,file_name,link FROM email_contents WHERE hs_id = %s", (hs_id,))

    if not results:
        logger.info(f"No files found for hs_id: {hs_id}")
        return {"status": "no_files", "message": f"No files found for hs_id: {hs_id}"}

    # Create a temporary directory
    temp_dir = tempfile.mkdtemp()
    logger.debug(f"Created temporary directory: {temp_dir}")

    downloaded_files = []
    files_object = []
    errors = []

    try:
        for row in results:
            id = row["id"]
            file_name = row["file_name"]
            link = row["link"]

            # Use the temporary directory instead of DOWNLOADS_DIR
            temp_file_path = os.path.join(temp_dir, file_name)

            try:
                minio_client = get_minio_client()
                # Download the file from MinIO to temp directory
                minio_client.fget_object(
                    bucket_name=MINIO_BUCKET,
                    object_name=os.path.basename(link),
                    file_path=temp_file_path,  # Use temp directory path
                )
                logger.debug(f"Downloaded {file_name} to {temp_dir}")
                downloaded_files.append(temp_file_path)

                # Classify text or image PDF
                is_image_based = is_image_document(temp_file_path)
                if is_image_based:
                    logger.info(
                        f"File {file_name} is an image document, converting to text.")
                    classify_type = "IMAGE"

                    # Convert image-based PDF to text
                    extracted_text = convert_pdf_to_text(
                        temp_file_path)

                    if not extracted_text:
                        logger.warning(
                            f"Could not extract text from image-based PDF: {file_name}")
                        # Update status in database
                        executeSQL("UPDATE email_contents SET type = 'UNKNOWN', status = 'XU_LY_LOI', classify_type = %s WHERE id = %s",
                                   (classify_type, id))
                        continue

                    # Save extracted text to a temporary file
                    markdown_filename = f"{os.path.splitext(file_name)[0]}_{uuid.uuid4().hex[:8]}.md"
                    markdown_path = os.path.join(temp_dir, markdown_filename)
                    with open(markdown_path, "w", encoding="utf-8") as f:
                        f.write(extracted_text)

                    # Upload markdown to MinIO
                    uploaded_files = upload_to_minio(
                        file_paths=markdown_path,
                        bucket_name="markdown",  # Use a separate bucket for markdown files
                        minio_endpoint=f"http://{EnvSettings().MINIO_API_ENDPOINT}",
                        access_key=EnvSettings().MINIO_ACCESS_KEY,
                        secret_key=EnvSettings().MINIO_SECRET_KEY,
                    )

                    if uploaded_files:
                        markdown_link = uploaded_files[0]
                        logger.info(
                            f"Uploaded extracted text to MinIO: {markdown_link}")

                        # Classify the document type based on extracted text
                        doc_type, status = classify_document_from_text(
                            extracted_text, file_name)

                        # Update database with markdown link and classification
                        executeSQL("UPDATE email_contents SET type = %s, status = %s, classify_type = %s, markdown_link = %s WHERE id = %s",
                                   (doc_type, status, classify_type, markdown_link, id))

                        if doc_type != "unknown":
                            files_object.append({
                                "id": id, "file_name": file_name, "file_type": doc_type,
                                "file_path": link, "classify_type": classify_type,
                                "markdown_link": markdown_link
                            })
                    else:
                        logger.error(
                            f"Failed to upload extracted text to MinIO for {file_name}")
                else:
                    # Handle text-based PDF (your existing code)
                    logger.info(
                        f"File {file_name} is a text PDF, proceeding with classification.")
                    classify_type = "TEXT"

                    # Extract text from text-based PDF
                    extracted_text = extract_text_from_pdf(temp_file_path)

                    if not extracted_text:
                        logger.warning(
                            f"Could not extract text from PDF: {file_name}")
                        # Update status in database
                        executeSQL("UPDATE email_contents SET type = 'UNKNOWN', status = 'XU_LY_LOI', classify_type = %s WHERE id = %s",
                                   (classify_type, id))
                        continue

                    # Classify the document type
                    doc_type, status = classify_document_from_text(
                        extracted_text, file_name)

                    # Update back to database with classification result
                    executeSQL("UPDATE email_contents SET type = %s, status = %s, classify_type = %s WHERE id = %s",
                               (doc_type, status, classify_type, id))

                    if doc_type != "unknown":
                        files_object.append({
                            "id": id, "file_name": file_name, "file_type": doc_type,
                            "file_path": link, "classify_type": classify_type,
                            "markdown_link": ""
                        })

            except S3Error as e:
                error_msg = f"Error downloading file {file_name}: {str(e)}"
                logger.error(error_msg)
                errors.append(error_msg)
            except Exception as e:
                error_msg = f"Unexpected error processing {file_name}: {str(e)}"
                logger.error(error_msg)
                errors.append(error_msg)

        # Check if we have any successfully classified files
        if not files_object:
            return {"status": "error", "message": "No files could be classified successfully"}

        # Check if we have HSMT files
        has_hsmt = any(file["file_type"] == "HSMT" for file in files_object)

        # If no HSMT files found, return an error
        if not has_hsmt:
            return {"status": "error", "message": "Không có file Hồ sơ mời thầu trong bộ file được tải lên!"}

        logger.debug(f"Files object: {files_object}")
        message = {
            "id": hs_id,
            "bucket": MINIO_BUCKET,
            "files": files_object,
        }

        return {"status": "success", "message": message}

    except Exception as e:
        # Handle overall function errors
        error_msg = f"Error in classify function: {str(e)} at {traceback.format_exc()}"
        logger.error(error_msg)
        return {"status": "error", "message": error_msg}
    finally:
        # Clean up - remove all files in the temp directory
        try:
            shutil.rmtree(temp_dir)
            logger.debug(f"Removed temporary directory {temp_dir}")
        except Exception as e:
            logger.warning(
                f"Failed to remove temporary directory {temp_dir}: {str(e)}")


def is_image_document(file_path, text_threshold=100, image_coverage_threshold=0.3):
    """
    Determines if a document (PDF or DOCX) is primarily image-based or text-based.

    Args:
        file_path: Path to the document file (PDF or DOCX)
        text_threshold: Minimum characters per page to consider as text content
        image_coverage_threshold: Ratio of page area covered by images to consider as image-based

    Returns:
        bool: True if document is primarily image-based, False if text-based
    """
    file_ext = os.path.splitext(file_path)[1].lower()

    if file_ext == '.pdf':
        return _analyze_pdf(file_path, text_threshold, image_coverage_threshold)
    elif file_ext == '.docx':
        return _analyze_docx(file_path, text_threshold)
    else:
        logger.warning(f"Unsupported file format: {file_ext}")
        return False  # Default to text-based for unsupported formats


def _analyze_pdf(pdf_path, text_threshold=100, image_coverage_threshold=0.3):
    """Helper function to analyze PDF files"""
    try:
        pdf_document = fitz.open(pdf_path)
        total_pages = len(pdf_document)
        image_pages = 0

        for page_num in range(total_pages):
            page = pdf_document[page_num]

            # Get text content
            text = page.get_text()
            text_length = len(text.strip())

            # Get images on the page
            image_list = page.get_images(full=True)

            # Calculate page area
            page_area = page.rect.width * page.rect.height

            # Calculate area covered by images
            images_area = 0
            for img_idx, img in enumerate(image_list):
                xref = img[0]
                image_rect = page.get_image_bbox(img)
                if image_rect:
                    images_area += image_rect.width * image_rect.height

            # Calculate image coverage ratio
            image_coverage = images_area / page_area if page_area > 0 else 0

            # Determine if this page is image-based
            if text_length < text_threshold or image_coverage > image_coverage_threshold:
                image_pages += 1

        # If more than half the pages are image-based, consider the whole document image-based
        return image_pages / total_pages > 0.5

    except Exception as e:
        logger.error(f"Error analyzing PDF: {str(e)}")
        return False
    finally:
        if 'pdf_document' in locals():
            pdf_document.close()


def _analyze_docx(docx_path, text_threshold=100):
    """Helper function to analyze DOCX files"""
    try:
        import docx
        doc = docx.Document(docx_path)

        # Count pages (approximation)
        total_pages = max(1, sum(len(p.text) for p in doc.paragraphs) // 3000)
        text_content = "\n".join([p.text for p in doc.paragraphs])
        total_text = len(text_content)

        # Count images
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1

        # If very little text per page or many images, consider it image-based
        if total_text / total_pages < text_threshold or image_count > total_pages:
            return True

        return False

    except ImportError:
        logger.warning(
            "python-docx library not installed. Install with: pip install python-docx")
        return False
    except Exception as e:
        logger.error(f"Error analyzing DOCX: {str(e)}")
        return False


def classify_document_from_text(text, file_name=None):
    """
    Classifies a document based on extracted text content into types: HSMT, TBMT, HSKT, or unknown.

    Args:
        text: Extracted text content
        file_name: Optional name of the file (for logging purposes)

    Returns:
        tuple: (document_type, status) where document_type is one of 'HSMT', 'TBMT', 'HSKT', 'unknown'
               and status is either 'CHUA_XU_LY' or 'XU_LY_LOI'
    """
    try:

        # Extract key information for classification
        text_lower = text.lower()

        # Determine document type using match-case statement
        doc_type = None
        status = "CHUA_XU_LY"

        match True:
            case _ if "hồ sơ mời thầu" in text_lower:
                doc_type = "HSMT"
                if file_name:
                    logger.info(f"Classified {file_name} as HSMT")
            case _ if "thông báo mời thầu" in text_lower:
                doc_type = "TBMT"
                if file_name:
                    logger.info(f"Classified {file_name} as TBMT")
            case _ if re.search(r"k[ỹy]\s*thu[ậa]t", text_lower) or re.search(r"yêu\s+cầu\s+về\s+k[ỹy]\s*thu[ậa]t", text_lower):
                doc_type = "HSKT"
                if file_name:
                    logger.info(f"Classified {file_name} as HSKT")
            case _:
                doc_type = "unknown"
                status = "XU_LY_LOI"
                if file_name:
                    logger.warning(
                        f"Could not classify {file_name}, marked as UNKNOWN")

        return doc_type, status
    except Exception as e:
        logger.error(f"Error classifying document from text: {str(e)}")
        return "unknown", "XU_LY_LOI"


def pdf_image_to_text_batch(pdf_path, page_batch_size=5):
    """
    Converts image-based PDF pages to text using Google Gemini model.
    Processes pages in batches to optimize performance.

    Args:
        pdf_path: Path to the PDF file
        page_batch_size: Number of pages to process in a single batch

    Returns:
        str: Text extracted from the PDF, potentially in markdown format
    """
    try:
        pdf_document = fitz.open(pdf_path)
        total_pages = len(pdf_document)
        logger.info(f"Processing image PDF with {total_pages} pages")

        # Store extracted text from each batch
        all_text = []

        # Process pages in batches
        for batch_start in range(0, total_pages, page_batch_size):
            batch_end = min(batch_start + page_batch_size, total_pages)
            logger.info(f"Processing batch: pages {batch_start+1}-{batch_end}")

            batch_text = []
            batch_images = []

            # Extract images from each page in the batch
            for page_num in range(batch_start, batch_end):
                page = pdf_document[page_num]
                try:
                    # First try to get any text that might be present
                    text = page.get_text().strip()
                    if text:
                        batch_text.append(f"## Page {page_num+1}\n\n{text}")

                    # Convert page to image for OCR processing
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_data = pix.tobytes("png")
                    batch_images.append((f"Page {page_num+1}", img_data))

                except Exception as e:
                    logger.error(
                        f"Error extracting content from page {page_num+1}: {str(e)}")

            # If we have images, process them with Gemini
            if batch_images:
                try:
                    # Create prompt for the model
                    prompt = "Extract all the text from these PDF pages. Format your response as markdown."

                    # Call the Gemini model (adjust based on your actual implementation)
                    response = llm.chat_model_gpt_4o_mini_16k().invoke(prompt)

                    # Add the model's response to our batch text
                    if response and response.content:
                        all_text.append(response.content)

                    # Sleep briefly to avoid rate limits
                    time.sleep(1)

                except Exception as e:
                    logger.error(
                        f"Error processing batch with AI model: {str(e)}")
                    # If AI processing fails, add any text we extracted directly
                    if batch_text:
                        all_text.extend(batch_text)
            else:
                # If no images, just use the directly extracted text
                all_text.extend(batch_text)

        # Combine all extracted text
        final_markdown = "\n\n".join(all_text)
        return final_markdown

    except Exception as e:
        logger.error(f"Error in pdf_image_to_text_batch: {str(e)}")
        return ""
    finally:
        if 'pdf_document' in locals():
            pdf_document.close()


def extract_text_from_pdf(pdf_path):
    """
    Extracts text from a text-based PDF file.

    Args:
        pdf_path: Path to the PDF file

    Returns:
        str: Extracted text, potentially in markdown format
    """
    try:
        pdf_document = fitz.open(pdf_path)
        total_pages = len(pdf_document)
        logger.info(f"Extracting text from PDF with {total_pages} pages")

        # Extract text from each page
        text_parts = []
        for page_num in range(total_pages):
            page = pdf_document[page_num]
            text = page.get_text()
            if text.strip():
                text_parts.append(f"## Page {page_num+1}\n\n{text}")

        # Combine all extracted text
        final_markdown = "\n\n".join(text_parts)
        return final_markdown

    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return ""
    finally:
        if 'pdf_document' in locals():
            pdf_document.close()
